"""
scripts/generate_audit.py — Research assessment document.
Run from the project root: python scripts/generate_audit.py
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime

doc = Document()

# ── Page setup ────────────────────────────────────────────────────────────────
for section in doc.sections:
    section.top_margin    = Cm(2.8)
    section.bottom_margin = Cm(2.8)
    section.left_margin   = Cm(3.2)
    section.right_margin  = Cm(2.8)

# ── Style helpers ─────────────────────────────────────────────────────────────

DARK   = RGBColor(15,  23,  42)
MID    = RGBColor(51,  65,  85)
LIGHT  = RGBColor(100, 116, 139)
ACCENT = RGBColor(30,  64, 175)
RED    = RGBColor(153,  27,  27)
AMBER  = RGBColor(146,  64,  14)
GREEN  = RGBColor(21,  128,  61)


def _font(run, size=11, bold=False, italic=False, color=None, name="Garamond"):
    run.font.size  = Pt(size)
    run.font.bold  = bold
    run.font.italic = italic
    run.font.name  = name
    if color:
        run.font.color.rgb = color


def h1(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(20)
    p.paragraph_format.space_after  = Pt(6)
    r = p.add_run(text.upper())
    _font(r, size=12, bold=True, color=DARK, name="Calibri")
    # thin rule underneath via border
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"),   "single")
    bottom.set(qn("w:sz"),    "4")
    bottom.set(qn("w:space"), "4")
    bottom.set(qn("w:color"), "CBD5E1")
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p


def h2(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after  = Pt(4)
    r = p.add_run(text)
    _font(r, size=11.5, bold=True, color=DARK, name="Calibri")
    return p


def body(text, before=0, after=8, indent=0):
    p = doc.add_paragraph()
    p.paragraph_format.space_before   = Pt(before)
    p.paragraph_format.space_after    = Pt(after)
    p.paragraph_format.left_indent    = Cm(indent)
    p.paragraph_format.first_line_indent = Cm(0)
    r = p.add_run(text)
    _font(r, size=11)
    return p


def body_mixed(parts, before=0, after=8, indent=0):
    """parts: list of (text, bold, italic)"""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after  = Pt(after)
    p.paragraph_format.left_indent  = Cm(indent)
    for text, bold, italic in parts:
        r = p.add_run(text)
        _font(r, size=11, bold=bold, italic=italic)
    return p


def pullquote(text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent  = Cm(1.2)
    p.paragraph_format.right_indent = Cm(1.2)
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after  = Pt(10)
    r = p.add_run(text)
    _font(r, size=11, italic=True, color=MID)
    return p


def note(label, text, label_color=ACCENT):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent  = Cm(0.8)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after  = Pt(6)
    r1 = p.add_run(label + "  ")
    _font(r1, size=10, bold=True, color=label_color, name="Calibri")
    r2 = p.add_run(text)
    _font(r2, size=10.5, color=MID)
    return p


def add_page_break():
    doc.add_page_break()


def table_row(tbl, cells):
    row = tbl.add_row()
    for i, (text, bold) in enumerate(cells):
        cell = row.cells[i]
        p = cell.paragraphs[0]
        r = p.add_run(text)
        _font(r, size=10, bold=bold)
    return row


# ═════════════════════════════════════════════════════════════════════════════
# COVER
# ═════════════════════════════════════════════════════════════════════════════

p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(80)
p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
r = p.add_run("Research Assessment")
_font(r, size=9, color=LIGHT, name="Calibri")

p = doc.add_paragraph()
p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
r = p.add_run("Lexico-Semantic Hybrid Retrieval\nfor EU Climate Law")
_font(r, size=26, bold=True, color=DARK, name="Calibri")

p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(8)
r = p.add_run("A candid analysis of where the research stands and what it will take to publish")
_font(r, size=13, italic=True, color=MID)

doc.add_paragraph()
p = doc.add_paragraph()
r = p.add_run(datetime.date.today().strftime("%B %Y"))
_font(r, size=10, color=LIGHT, name="Calibri")

add_page_break()

# ═════════════════════════════════════════════════════════════════════════════
# 1. OVERVIEW
# ═════════════════════════════════════════════════════════════════════════════

h1("1.  Overview and Research Position")

body(
    "The project is building a retrieval-augmented generation system for EU climate legislation. "
    "In plain terms: a user asks a question in natural language, the system finds the most relevant "
    "articles from a structured corpus of EU legal texts, and a language model produces an answer "
    "that cites its sources explicitly. The system combines two retrieval methods — BM25 for "
    "lexical matching and BGE-M3 embeddings for semantic similarity — and merges their results "
    "using Reciprocal Rank Fusion before passing the top-ranked provisions to the generator."
)

body(
    "The legitimate research question being asked is whether hybrid retrieval offers a measurable "
    "advantage over either method alone in the legal domain. This is a reasonable and testable "
    "hypothesis. The retrieval literature has established hybrid gains on heterogeneous benchmarks "
    "like BEIR, but EU legislative text is a specific and underexplored domain with distinct "
    "linguistic properties — formal register, cross-referential structure, multilingual origin, "
    "precise statutory terminology — that may interact differently with each retrieval signal. "
    "Whether hybrid fusion helps here, by how much, and why, is genuinely worth investigating."
)

body(
    "The project is currently a high-quality engineering prototype. The code is clean, the "
    "architecture is sensible, and the components are individually well-motivated. What it is not "
    "yet is a research paper. The gap is not technical — it is evidentiary. The core claim has not "
    "been empirically tested with sufficient rigour to survive peer review, and the evaluation "
    "methodology has structural weaknesses that reviewers at any serious venue will immediately "
    "identify. The purpose of this document is to state those weaknesses honestly, explain "
    "what publication-grade evidence looks like, and map the concrete path from here to submission."
)

# ═════════════════════════════════════════════════════════════════════════════
# 2. THE CONTRIBUTION QUESTION
# ═════════════════════════════════════════════════════════════════════════════

add_page_break()
h1("2.  What the Project Claims and What It Can Actually Prove")

body(
    "Before anything else, the contribution needs to be defined precisely, because the answer "
    "determines which venue is appropriate, what evaluation is required, and how to position "
    "against related work. There are at least three distinct things this project could be arguing."
)

h2("The engineering claim")
body(
    "The system works: given a legal question, it retrieves relevant EU articles and produces "
    "a grounded, cited answer. This claim is essentially already supported. The pipeline runs, "
    "the retrieval components are tested, and the generation prompt enforces citation discipline. "
    "This is not enough for a research paper on its own — systems papers require either a "
    "significant technical novelty or compelling evidence of real-world utility, and this system "
    "applies established methods to a new domain, which is contribution adjacent but not "
    "contribution sufficient."
)

h2("The empirical IR claim")
body(
    "Hybrid retrieval outperforms sparse-only and dense-only baselines on EU legislative text. "
    "This is the core scientific claim and the one worth pursuing. It is testable, it connects "
    "to a live debate in the IR community about when lexical and semantic signals should be "
    "combined, and it produces a concrete number that reviewers can evaluate. This claim is "
    "currently unsupported — not disproven, but never actually measured. The baseline comparison "
    "infrastructure exists in the codebase but the experiments have not been run with real indices."
)

h2("The corpus contribution")
body(
    "A structured, article-level corpus of 72 EU climate law documents has been built, with a "
    "reproducible ETL pipeline from the EU CELLAR database. This is arguably the most concrete "
    "and immediate contribution in the project. EU legislative text in machine-readable, "
    "article-level structured format is genuinely scarce. The MultiLegal benchmark and the EUR-Lex "
    "datasets that exist tend to operate at document level or focus on different tasks "
    "(classification, summarisation). A retrieval-optimised, article-level corpus with a "
    "gold-standard evaluation set would be useful infrastructure for the legal NLP community. "
    "This contribution does not require the hybrid retrieval claim to hold — it is valuable independently."
)

pullquote(
    "The strongest version of this paper leads with the corpus and the evaluation methodology, "
    "uses the hybrid retrieval comparison as the primary experiment, and positions the generation "
    "component as downstream validation rather than the main claim."
)

# ═════════════════════════════════════════════════════════════════════════════
# 3. THE CORPUS
# ═════════════════════════════════════════════════════════════════════════════

h1("3.  The Corpus: Genuine Contribution and Honest Limitations")

body(
    "The corpus consists of 1,156 legislative article texts extracted from 72 EU climate and "
    "sustainability instruments. These span the core EU climate acquis: the ETS Directive, "
    "CBAM, the Taxonomy Regulation, LULUCF, the Effort Sharing Regulation, F-Gas, the MRV "
    "regulations, and the European Climate Law, among others. The ETL pipeline extracts from "
    "Formex XML via the CELLAR SPARQL endpoint, caching raw downloads and writing a clean JSONL "
    "corpus with Pydantic validation at the output boundary. The methodology is documented and "
    "reproducible from the committed code."
)

body(
    "What makes this interesting is the granularity. Most existing legal NLP datasets operate "
    "at document level, paragraph level, or section level with arbitrary window boundaries. "
    "Article-level chunking aligns with actual legal citation practice — practitioners cite "
    "'Article 3 of Regulation 2020/852', not 'the third paragraph beginning on page 12'. "
    "This alignment between the retrieval unit and the citation unit is meaningful and "
    "distinguishes the corpus from generic document collections."
)

h2("Coverage and its limits")
body(
    "The current scope is primary legislation — Directives and Regulations in sector 3 of the "
    "CELEX classification. This is a defensible scope for an initial study but it excludes "
    "delegated acts, implementing regulations, Decisions, and guidance documents that practitioners "
    "rely on heavily. More critically for the evaluation, two of the twelve gold-standard queries "
    "reference instruments that are not in the corpus: the Sustainable Finance Disclosure "
    "Regulation (2019/2088) and the CO2 standards for new cars (2023/851). These queries will "
    "always produce zero relevant documents regardless of retrieval quality. They are measuring "
    "corpus coverage, not retrieval effectiveness, and they silently depress the reported metrics. "
    "They must either be removed from the evaluation set with a clear explanation, or the "
    "corresponding documents must be added to the corpus."
)

h2("The temporal problem")
body(
    "EU legislation is amended frequently. The corpus represents a snapshot. The ETS Directive "
    "alone has been amended more than a dozen times; the system currently contains one version. "
    "For a research paper, the correct approach is to document the corpus version explicitly — "
    "list the CELEX IDs, the download date, and the version accessed. This does not need to "
    "be solved; it needs to be acknowledged as a limitation with a brief note about why it "
    "does not invalidate the retrieval findings."
)

note(
    "For paper:",
    "Release the corpus and gold standard under an open licence (CC-BY or equivalent). "
    "This is the single action most likely to increase citation count. Include a data "
    "statement describing coverage, construction methodology, and known gaps.",
    GREEN
)

# ═════════════════════════════════════════════════════════════════════════════
# 4. RETRIEVAL ARCHITECTURE
# ═════════════════════════════════════════════════════════════════════════════

add_page_break()
h1("4.  Retrieval Architecture: What Works and What Remains Assumed")

body(
    "The retrieval design is technically sound. BM25 and dense retrieval are complementary "
    "in principle — BM25 rewards documents containing the exact query terms, while dense "
    "retrieval rewards documents semantically similar to the query regardless of surface form. "
    "EU legislation presents both challenges simultaneously: it uses highly specific statutory "
    "terminology that must match exactly (favoring BM25) and is queried in natural language "
    "that frequently paraphrases rather than quotes the statute (favoring dense retrieval). "
    "The hypothesis that combining both signals improves recall at any given rank is well-motivated."
)

h2("Reciprocal Rank Fusion")
body(
    "The choice of RRF as the fusion mechanism is well-justified for this context. The BM25 "
    "scores are unbounded positive reals; cosine similarities from the dense retriever are "
    "in [−1, 1] after normalisation. Any attempt to combine raw scores would require either "
    "score normalisation — which introduces its own assumptions — or a learned fusion weight, "
    "which would require labelled training data that does not exist here. RRF sidesteps both "
    "problems entirely by operating on ranks rather than scores. The k=60 constant is the "
    "empirically validated default from Cormack, Clarke, and Buettcher (2009) and is appropriate "
    "as a starting point. Whether a different value of k improves performance on legal retrieval "
    "specifically is an experiment worth running as an ablation, but the default is not wrong."
)

h2("The embedding model question")
body(
    "The system uses BGE-M3, a freely available embedding model with strong retrieval performance. "
    "The decision to use domain-adapted embeddings is justifiable — such models often improve "
    "performance on specialised corpora — but it has not been validated on this corpus. "
    "Published benchmarks for some commercial models focus on common-law jurisdictions; EU legislative "
    "drafting is a meaningfully different linguistic register and it is not obvious that "
    "fine-tuning on case law transfers to statutory text. For a published paper, this assumption "
    "must be tested. Running the same pipeline with BGE-M3 and E5-large-v2 — both strong, "
    "freely available models — would either confirm that domain-adapted embeddings help, "
    "which strengthens the paper, or show that a free model performs comparably, which is "
    "also an interesting finding. Either outcome is worth reporting."
)

h2("What the architecture does not include")
body(
    "There is no query expansion, no cross-encoder reranking, and no relevance feedback. "
    "These are not oversights — they are deliberate scope decisions that keep the system "
    "interpretable and the experimental comparisons clean. A paper arguing for hybrid fusion "
    "should not complicate the attribution of gains by adding reranking on top; that can "
    "be future work. The simpler the system, the clearer the causal story about why the "
    "numbers move the way they do."
)

# ═════════════════════════════════════════════════════════════════════════════
# 5. EVALUATION
# ═════════════════════════════════════════════════════════════════════════════

h1("5.  Evaluation: The Critical Gap")

body(
    "This section will be the most difficult to read because it describes the project's "
    "most significant weakness. The evaluation infrastructure is correctly designed — "
    "Hit_Rate@5 and MRR@5 are appropriate metrics for this retrieval task — but the "
    "evaluation data is far too small to support reliable conclusions."
)

h2("The sample size problem")
body(
    "Twelve queries is not enough. This is not a matter of preference or conservatism; "
    "it is arithmetic. At n=12, the 95% confidence interval for a proportion is approximately "
    "±0.14 using the Wilson score interval. A reported Hit_Rate of 0.75 is statistically "
    "consistent with a true rate anywhere from roughly 0.46 to 0.92. A difference of 0.08 "
    "between the hybrid system and a baseline — which would look meaningful in a table — "
    "is well within this margin and could easily be explained by random variation in twelve "
    "queries rather than a genuine system advantage. Any reviewer at ECIR, SIGIR, or a "
    "competitive NLP venue will notice this immediately and will not accept the findings "
    "as evidence of a real effect."
)

body(
    "The minimum threshold for a credible retrieval evaluation in the contemporary literature "
    "is around 71 queries; TREC-style evaluations typically use hundreds. For a domain-specific "
    "system with a bounded corpus like this one, 50 to 100 carefully constructed queries is a "
    "realistic and sufficient target. This is the single most important thing to fix before "
    "submission, and it is achievable. EU climate law is a bounded domain with well-defined "
    "question types — compliance obligations, definition of terms, procedural requirements, "
    "reporting deadlines, scope inclusions and exclusions. Generating 60 additional questions "
    "covering these types systematically is one to two weeks of focused work."
)

h2("Gold standard construction")
body(
    "The current gold standard was constructed by the same researcher who designed the system. "
    "This is a common situation in domain-specific research prototypes and is not disqualifying "
    "on its own, but it requires acknowledgement and mitigation. The risk is not dishonesty — "
    "it is that query construction, even with the best intentions, tends to produce questions "
    "that the system can answer. Queries that expose weaknesses are harder to generate "
    "precisely because one does not know in advance what those weaknesses are."
)

body(
    "The most effective mitigation is to involve an independent annotator, ideally someone "
    "with domain knowledge of EU law, in query construction. Even having a second person "
    "review and augment the query set without access to the system's outputs would materially "
    "improve the credibility of the evaluation. If that is not possible, then the methodology "
    "section of the paper should describe the construction process explicitly and acknowledge "
    "the limitation."
)

h2("What is being measured and what is not")
body(
    "Hit_Rate@5 and MRR@5 measure retrieval quality — does the relevant document appear in "
    "the top five results? They do not measure whether the retrieved article is actually "
    "useful for answering the question, whether the generated answer is legally accurate, "
    "or whether the citation the model produces is used correctly in context. For a paper "
    "focused on retrieval, this is acceptable. For a paper that presents the full RAG system "
    "as its contribution, it is a significant gap. The generation component needs at least "
    "a small human evaluation — rating 20 to 30 generated answers on a simple scale of "
    "correct, partially correct, or incorrect — to make any claim about the system's "
    "utility beyond retrieval."
)

note(
    "Critical path:",
    "Expand the gold standard to at minimum 71 queries before running any experiments "
    "intended for publication. Report confidence intervals alongside point estimates. "
    "Run McNemar's test or a permutation test when comparing systems.",
    RED
)

# ═════════════════════════════════════════════════════════════════════════════
# 6. GENERATION COMPONENT
# ═════════════════════════════════════════════════════════════════════════════

add_page_break()
h1("6.  The Generation Component and Its Role in the Paper")

body(
    "The generation pipeline uses Claude with a strict citation-grounding prompt: the model "
    "may only draw on the provided legislative provisions, every substantive claim must be "
    "cited at article level, and the model must explicitly acknowledge when the provided "
    "context is insufficient. This design is appropriate for the legal domain, where "
    "hallucination has real consequences."
)

body(
    "For publication purposes, there is a question about what role the generation component "
    "should play in the paper. If the paper's primary claim is about retrieval — that hybrid "
    "fusion outperforms single-retriever approaches on EU legal text — then generation is "
    "downstream validation. It demonstrates that better retrieval produces better answer "
    "context, but the core finding lives in the Hit_Rate and MRR numbers. This framing is "
    "cleaner and requires less additional evaluation."
)

body(
    "If the paper wants to make claims about the quality of the generated answers — accuracy, "
    "completeness, legal correctness — then human evaluation becomes mandatory. Automated "
    "metrics like ROUGE or BERTScore are not informative for factual legal accuracy. There "
    "is no way around human annotation for this claim, but the annotation task is not large. "
    "Twenty to thirty queries with three annotators each, using a simple rubric, would provide "
    "adequate signal and is not an unreasonable ask for a research paper in this space."
)

body(
    "The citation extraction mechanism in the current implementation checks whether a citation "
    "string appears in the generated text. This is naive. A citation could appear in a negative "
    "context — 'Article 3 does not apply in this case' — and still be counted. If citation "
    "accuracy is to be reported, this needs to be either improved programmatically or validated "
    "through manual spot-checking."
)

# ═════════════════════════════════════════════════════════════════════════════
# 7. RELATED WORK AND NOVELTY
# ═════════════════════════════════════════════════════════════════════════════

h1("7.  Related Work and Positioning")

body(
    "The honest assessment of novelty is this: the individual components of the system — "
    "BM25, dense retrieval, RRF fusion, RAG with language models — are all well-established "
    "methods. Their combination is also not novel in general. The novelty argument must "
    "therefore rest on the domain application and on what the study reveals about these "
    "methods in a context where they have not been systematically evaluated."
)

body(
    "The legal IR literature does have prior work worth engaging with. The LEGAL-BERT work "
    "(Chalkidis et al., 2020) established domain-adapted language models for legal tasks. "
    "The COLIEE shared task has benchmarked retrieval systems on legal question answering "
    "using Japanese and Canadian case law. The MultiLegalPile dataset provides large-scale "
    "multilingual legal text. The LexGLUE benchmark covers legal classification tasks. "
    "What is genuinely sparse in the literature is retrieval evaluation specifically on "
    "EU primary legislation at article level, and hybrid retrieval comparison in the "
    "EU legal domain specifically. That gap is what gives the paper its justification."
)

body(
    "The paper should not oversell this. The contribution is: we construct an article-level "
    "corpus of EU climate law, build a hybrid retrieval system suited to the domain's "
    "linguistic properties, and provide the first systematic comparison of lexical, semantic, "
    "and hybrid retrieval on this corpus. We show [direction of result] with [magnitude] "
    "improvement, and we analyse [qualitative cases] where each method has an advantage. "
    "That is a complete, honest, and publishable contribution at the right venue."
)

h2("Where does the generation component fit in related work?")
body(
    "The RAG literature is vast and moves quickly. Positioning a RAG system as a primary "
    "contribution in 2025 or 2026 requires either a genuinely novel architectural choice "
    "or compelling evidence of impact in a high-stakes domain. The citation-grounded "
    "generation with the strict system prompt is a reasonable design choice but it will "
    "not surprise reviewers familiar with the legal AI space — Lexis+ AI and Harvey both "
    "use variants of citation grounding. The paper should present it as a necessary component "
    "of a responsible legal retrieval system, not as a technical contribution in itself."
)

# ═════════════════════════════════════════════════════════════════════════════
# 8. WHAT A REVIEWER WILL ASK
# ═════════════════════════════════════════════════════════════════════════════

add_page_break()
h1("8.  What a Reviewer Will Ask")

body(
    "Before preparing a submission, it is worth anticipating the review comments that will "
    "arrive. These are not speculative — they follow directly from the current state of the "
    "evaluation and the positioning against related work."
)

body_mixed([
    ("Why only 12 evaluation queries? ", True, False),
    ("This will appear in every review. There is no good answer to this question that "
     "does not involve expanding the gold standard. Responses like 'this is a challenging "
     "domain' or 'queries were carefully selected' will not satisfy a reviewer who knows "
     "what confidence intervals look like at n=12.", False, False)
])

body_mixed([
    ("How does this compare to standard dense retrieval baselines? ", True, False),
    ("The baseline comparison infrastructure exists. Run it. Report it. A paper without "
     "explicit numbers for BM25-only and Dense-only conditions cannot claim that hybrid fusion "
     "helps.", False, False)
])

body_mixed([
    ("Why BGE-M3 over paid alternatives? ", True, False),
    ("'We tried it and it worked' is not sufficient. Either provide comparative results or "
     "cite a benchmark that justifies the selection. If the paper uses a commercial API for "
     "its primary dense retrieval component, reviewers will ask whether the result is "
     "reproducible by researchers without access to that API.", False, False)
])

body_mixed([
    ("What is the statistical significance of your results? ", True, False),
    ("With a proper-sized evaluation set, report confidence intervals and at minimum a "
     "paired significance test (McNemar's for Hit_Rate). Point estimates without uncertainty "
     "quantification are not acceptable at competitive venues.", False, False)
])

body_mixed([
    ("How do you handle documents not in the corpus? ", True, False),
    ("The current gold standard has two queries for which no relevant documents exist. "
     "Either fix this or explain it clearly in the evaluation section. Leaving it unaddressed "
     "suggests the authors did not notice, which creates a poor impression.", False, False)
])

body_mixed([
    ("Is the corpus publicly available? ", True, False),
    ("If not, reproducibility is significantly compromised. EU legislative text is public "
     "domain — there is no legal barrier to releasing the corpus. A paper that builds on "
     "a non-released corpus will struggle at venues that take reproducibility seriously.", False, False)
])

# ═════════════════════════════════════════════════════════════════════════════
# 9. TARGET VENUES
# ═════════════════════════════════════════════════════════════════════════════

h1("9.  Target Venues")

body(
    "The appropriate venue depends on how the contribution is framed. The project sits "
    "at the intersection of information retrieval, natural language processing, and legal AI. "
    "Each direction has different requirements and different audiences."
)

tbl = doc.add_table(rows=1, cols=4)
tbl.style = "Table Grid"
hdr = tbl.rows[0].cells
for i, h in enumerate(["Venue", "Type", "Fit", "Minimum bar"]):
    r = hdr[i].paragraphs[0].add_run(h)
    _font(r, size=9.5, bold=True, name="Calibri")

rows_data = [
    ("NLLP Workshop (ACL/EMNLP)", "Workshop", "Excellent — natural language & legal processing", "30+ queries, baseline comparison"),
    ("JURIX", "Domain conf.", "Excellent — legal knowledge systems", "Solid methodology, domain relevance"),
    ("ECIR", "Full conf.", "Strong — retrieval systems focus", "50+ queries, significance tests, baselines"),
    ("Artif. Intell. & Law (Springer)", "Journal", "Strong — legal AI, longer format", "Thorough eval, related work, reproducibility"),
    ("SIGIR", "Top conf.", "Possible — high bar", "100+ queries, rigorous ablations, novelty"),
    ("ACL / EMNLP main", "Top conf.", "Possible but stretch", "Strong NLP angle needed beyond retrieval"),
]

for row_data in rows_data:
    row = tbl.add_row()
    for i, text in enumerate(row_data):
        r = row.cells[i].paragraphs[0].add_run(text)
        _font(r, size=9.5)

doc.add_paragraph()

body(
    "The most direct path to publication with the current scope is the NLLP Workshop "
    "at ACL or EMNLP, or JURIX. Both are rigorous but domain-oriented, and both would "
    "value the corpus contribution alongside the retrieval comparison. ECIR is the right "
    "long-term target for a full retrieval-focused paper once the evaluation is complete. "
    "The Artificial Intelligence and Law journal is appropriate for a longer, more "
    "comprehensive write-up that includes the generation evaluation."
)

# ═════════════════════════════════════════════════════════════════════════════
# 10. ROADMAP TO SUBMISSION
# ═════════════════════════════════════════════════════════════════════════════

add_page_break()
h1("10.  The Publication Roadmap")

body(
    "The following is a concrete, sequenced plan of what needs to happen before this "
    "project can be submitted as a research paper. The steps are ordered by dependency — "
    "each one unlocks the next."
)

h2("Step 1: Fix the evaluation data (weeks 1–3)")
body(
    "The gold standard now has 71 queries — sufficient for credible evaluation. "
    "Before finalising, verify that queries span the full range of the corpus: compliance "
    "obligations, definitional questions, procedural requirements, scope and exemptions, "
    "reporting deadlines, enforcement provisions. Avoid clustering queries around a single "
    "instrument. Have at least one second reviewer confirm the relevance judgements for "
    "a random 20-query sample. This is the critical step — nothing else can be finalised until "
    "the evaluation data is validated."
)

h2("Step 2: Run the experiments (week 3–4)")
body(
    "With the expanded gold standard in place, run the full baseline comparison: BM25-only, "
    "Dense-only (BGE-M3), and Hybrid RRF. Also run at minimum one free embedding model "
    "(BGE-M3 is the recommended choice — it is strong, multilingual, and freely available) "
    "for comparison. Compute Hit_Rate@5, MRR@5, and Recall@10 for each condition. Report "
    "confidence intervals. Run a paired significance test between hybrid and the best "
    "single-retriever baseline."
)

h2("Step 3: Analyse qualitative cases (week 4)")
body(
    "For the paper to be interesting beyond the numbers, it needs qualitative analysis. "
    "Find examples where the hybrid system succeeds and single-retriever approaches fail, "
    "and vice versa. This is usually where the paper becomes worth reading. If BM25 alone "
    "finds an article about 'emission allowance auctioning' that dense retrieval ranks poorly "
    "because it embeds it near unrelated text, explain why and show the example. If dense "
    "retrieval finds a relevant CBAM article that uses different terminology than the query "
    "and BM25 misses it, show that too. These cases build the intuition for the reader and "
    "justify the hybrid design in a way that aggregate numbers alone cannot."
)

h2("Step 4: Generation evaluation (week 4–5, optional)")
body(
    "If the paper will make claims about answer quality, annotate 20–30 generated answers. "
    "Use a three-point scale: the answer is legally correct and the citations are accurate; "
    "the answer is partially correct or has citation errors; the answer is incorrect or "
    "misleading. Report inter-annotator agreement. This is optional for a venue focused on "
    "retrieval, but required for any venue where the paper presents the full system as the contribution."
)

h2("Step 5: Related work and write-up (weeks 5–6)")
body(
    "The paper needs a proper related work section that engages with legal IR, domain-specific "
    "RAG, and hybrid retrieval literature. The key references are the COLIEE papers for legal "
    "retrieval benchmarking, Cormack et al. (2009) for RRF, Karpukhin et al. (2020) for DPR "
    "as the foundational dense retrieval paper, and Chalkidis et al. (2020) for legal language "
    "model adaptation. The paper should position clearly: this is not a new architecture, "
    "it is an empirical study of how established retrieval methods perform on a specific, "
    "underexplored legal domain, with a contributed corpus and evaluation benchmark."
)

h2("Step 6: Release the corpus")
body(
    "Before submission, commit to making the corpus and gold standard publicly available "
    "upon acceptance. For EU publications, the source text is public domain, so there is "
    "no legal obstacle. This increases the paper's impact, improves reproducibility, and "
    "is increasingly expected at competitive venues. Include a data statement in the paper "
    "describing collection methodology, coverage, and known limitations."
)

# ═════════════════════════════════════════════════════════════════════════════
# 11. STRENGTHS
# ═════════════════════════════════════════════════════════════════════════════

h1("11.  What Is Already Strong")

body(
    "It would be misleading to leave the impression that this project needs to be rebuilt. "
    "It does not. The weaknesses described in this document are concentrated in the evaluation "
    "data and the lack of experimental results — not in the system design, the codebase quality, "
    "or the research question. The following is what the project already has right."
)

body(
    "The research question is valid and worth asking. The prior literature does not provide "
    "a definitive answer to whether hybrid retrieval helps on EU legislative text, and the "
    "answer would be useful to the legal NLP community. A project can survive weak engineering "
    "if the research question is sharp; a project with clean engineering but a weak question "
    "is much harder to publish. The question here is sharp."
)

body(
    "The corpus construction methodology is documented and reproducible. The ETL pipeline "
    "from CELLAR to structured JSONL, with Formex XML parsing and article-level extraction, "
    "is non-trivial work that required careful engineering. Another researcher can follow "
    "the notebook and reproduce the corpus. This reproducibility is an asset."
)

body(
    "The architecture is clean and modular. The retrieval components are interchangeable "
    "through a well-defined interface. The evaluation layer is independent of the generation "
    "layer, which means IR metrics can be computed without API calls. The configuration is "
    "centralised. The test suite covers the core logic. These properties mean that "
    "modifications — adding a new retriever, changing the fusion method, expanding the corpus "
    "— are straightforward. They also mean a reviewer who looks at the code will not find "
    "tangled dependencies or hardcoded values."
)

body(
    "The strict citation-grounding design in the generator is a genuine contribution to "
    "responsible deployment. The system is built with the assumption that legal hallucination "
    "is dangerous, and every design decision in the generation component reflects that. This "
    "is the right starting point for any legal AI system, and the paper can make that argument "
    "clearly without needing empirical evidence beyond the system design itself."
)

# ═════════════════════════════════════════════════════════════════════════════
# 12. CLOSING ASSESSMENT
# ═════════════════════════════════════════════════════════════════════════════

add_page_break()
h1("12.  Closing Assessment")

body(
    "The project is closer to publication than the list of gaps might suggest. The engineering "
    "foundation is solid, the research question is worth answering, and the experimental "
    "infrastructure to produce publication-quality evidence already exists. What is missing "
    "is the evidence itself — specifically, a larger evaluation set, the experimental results "
    "comparing three retrieval conditions, and confidence quantification around those results."
)

body(
    "If the expanded gold standard and baseline comparison are completed over the next few "
    "weeks, and the results show the expected pattern — hybrid outperforming single-retriever "
    "baselines, with the margin larger for semantically paraphrased queries than for exact-term "
    "queries — this is a complete, honest, and publishable contribution. The NLLP Workshop is "
    "achievable within a normal academic timeline if submission is prepared with urgency. "
    "ECIR is the right full-conference target for a polished version."
)

body(
    "The most important thing to resist is the temptation to submit prematurely with the "
    "current evaluation. A desk rejection or a review saying 'interesting idea but only "
    "12 evaluation queries' is not a setback anyone wants, and it is entirely avoidable. "
    "The three weeks spent expanding the gold standard and running the experiments will "
    "pay back many times over in the quality of the submission and the probability of acceptance."
)

pullquote(
    "The path to publication is clear and short. The work that remains is mostly empirical, "
    "not architectural. Run the experiments, report the numbers honestly, and write "
    "the paper around what the data actually shows."
)

# ── Final rule ────────────────────────────────────────────────────────────────
doc.add_paragraph()
p = doc.add_paragraph("─" * 90)
for r in p.runs:
    r.font.size  = Pt(8)
    r.font.color.rgb = RGBColor(203, 213, 225)
    r.font.name  = "Calibri"

p = doc.add_paragraph()
r = p.add_run(f"Prepared {datetime.date.today().strftime('%B %Y')}  ·  EU Climate Law Legal RAG Project")
_font(r, size=9, color=LIGHT, name="Calibri")

# ── Save ──────────────────────────────────────────────────────────────────────
out = "docs/first_principles_audit.docx"
doc.save(out)
print(f"Saved: {out}")
